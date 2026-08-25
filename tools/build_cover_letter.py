"""Build a cover letter DOCX + PDF from a markdown draft.

Usage: python3 tools/build_cover_letter.py <cover-letter.md> <output-dir>

The markdown format: first line `# <Company> - <Role>` becomes the subject,
blank-line-separated paragraphs become the body, and a trailing contact block
containing your email is dropped (the header renders identity instead).
Identity comes from `config/user-profile.json`.
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from user_config import contact_line, load_profile

PROFILE = load_profile()
IDENTITY = PROFILE["identity"]
NAME = IDENTITY["name"].upper()
ROLE = IDENTITY["headline_role"]
EMAIL = IDENTITY["email"]
CONTACT_PLAIN = contact_line(PROFILE)
CONTACT_PDF = CONTACT_PLAIN.replace(
    IDENTITY["linkedin"],
    f'<link href="{IDENTITY["linkedin_url"]}" color="#1F6FEB">{IDENTITY["linkedin"]}</link>',
)
BASENAME = PROFILE["files"]["cover_letter_basename"]

INK = HexColor("#000000")


def parse(md_path):
    """Return (subject, body_paragraphs) from the cover-letter markdown."""
    lines = Path(md_path).read_text().splitlines()
    subject = lines[0].lstrip("# ").strip()
    body, buf = [], []
    for line in lines[1:]:
        if line.strip():
            buf.append(line.strip())
        elif buf:
            body.append(" ".join(buf))
            buf = []
    if buf:
        body.append(" ".join(buf))
    # Drop the trailing contact line - it is rendered in the header instead.
    if body and EMAIL in body[-1]:
        body.pop()
    return subject, body


def esc(text):
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def build_pdf(out, subject, body):
    doc = SimpleDocTemplate(
        str(out), pagesize=letter,
        leftMargin=60, rightMargin=60, topMargin=52, bottomMargin=44,
        title=BASENAME, author=IDENTITY["name"],
    )
    st_name = ParagraphStyle("n", fontName="Helvetica-Bold", fontSize=20, leading=24, textColor=INK)
    st_role = ParagraphStyle("r", fontName="Helvetica", fontSize=11, leading=14, textColor=INK, spaceBefore=2)
    st_contact = ParagraphStyle("c", fontName="Helvetica", fontSize=9, leading=12, textColor=INK, spaceBefore=2)
    st_subj = ParagraphStyle("s", fontName="Helvetica-Bold", fontSize=10, leading=13, textColor=INK, spaceBefore=12)
    st_body = ParagraphStyle("b", fontName="Helvetica", fontSize=9.8, leading=12.9, textColor=INK, spaceBefore=8.5)

    flow = [
        Paragraph(NAME, st_name),
        Paragraph(esc(ROLE), st_role),
        Paragraph(CONTACT_PDF, st_contact),
        Paragraph("Re: " + esc(subject), st_subj),
        Spacer(1, 4),
    ]
    for para in body:
        flow.append(Paragraph(esc(para), st_body))
    doc.build(flow)


def build_docx(out, subject, body):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Helvetica"
    style.font.size = Pt(10.5)

    def para(text, size=10.5, bold=False, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)
        return p

    para(NAME, size=20, bold=True, space_after=0)
    para(ROLE, size=11, space_after=0)
    para(CONTACT_PLAIN, size=9, space_after=0)
    para("Re: " + subject, bold=True, space_before=14)
    for text in body:
        para(text, space_before=5, space_after=5)
    doc.save(str(out))


def main():
    md, out_dir = Path(sys.argv[1]), Path(sys.argv[2])
    subject, body = parse(md)
    out_dir.mkdir(parents=True, exist_ok=True)
    build_pdf(out_dir / f"{BASENAME}.pdf", subject, body)
    build_docx(out_dir / f"{BASENAME}.docx", subject, body)
    print(f"{len(body)} paragraphs -> {out_dir}")


if __name__ == "__main__":
    main()
